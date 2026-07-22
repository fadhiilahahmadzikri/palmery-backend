import io
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Any, Optional
from .interfaces import BaseExporter
from .templates.header_config import ReportHeaderConfig

class WordExporter(BaseExporter):
    def generate(self, records: List[Any], period_label: Optional[str] = None) -> io.BytesIO:
        document = Document()
        
        # Set landscape
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        
        header_data = ReportHeaderConfig.get_dynamic_header_data()
        
        # Header
        h1 = document.add_heading(header_data["company_name"], 0)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = document.add_paragraph()
        p.add_run(header_data["report_title"]).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if period_label:
            p_period = document.add_paragraph(f"Periode: {period_label}")
            p_period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_date = document.add_paragraph(f"Tanggal Cetak: {header_data['print_date']}")
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Table
        table = document.add_table(rows=1, cols=10)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = [
            "Tanggal", 
            "Pemanen", 
            "Lokasi", 
            "Janjang (jjg)", 
            "BJR (kg)", 
            "Bruto (kg)", 
            "Ptgn Brondolan (kg)", 
            "Mentah (jjg)", 
            "Denda Mentah (Rp)", 
            "Netto (kg)"
        ]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        for r in records:
            row_cells = table.add_row().cells
            date_str = r.harvest_date.strftime("%Y-%m-%d") if hasattr(r.harvest_date, 'strftime') else str(r.harvest_date)
            
            h_name = getattr(r, 'harvester_name', None)
            if not h_name and hasattr(r, 'harvester') and r.harvester:
                h_name = r.harvester.full_name
            if not h_name:
                h_name = str(getattr(r, 'harvester_id', '-'))

            loc_name = getattr(r, 'location_name', None)
            if not loc_name and hasattr(r, 'collection_point') and r.collection_point:
                loc_name = f"TPH {r.collection_point.point_number}"
            if not loc_name:
                loc_name = str(getattr(r, 'collection_point_id', '-'))

            gross = getattr(r, 'gross_tonnage_kg', 0.0)
            net = getattr(r, 'net_tonnage_kg', 0.0)
            loose_deduct = getattr(r, 'loose_fruit_deduction_kg', 0.0)
            fine_amount = getattr(r, 'fine_amount_rupiah', 0.0)
            unripe = getattr(r, 'unripe_bunch_count', 0)
            
            denda_str = f"Rp {fine_amount:,.0f}" if fine_amount > 0 else "-"

            row_cells[0].text = date_str
            row_cells[1].text = h_name
            row_cells[2].text = loc_name
            row_cells[3].text = str(getattr(r, 'valid_bunch_count', 0))
            row_cells[4].text = str(getattr(r, 'avg_bunch_weight_kg', 0))
            row_cells[5].text = f"{gross:,.1f}"
            row_cells[6].text = f"{loose_deduct:,.1f}"
            row_cells[7].text = str(unripe)
            row_cells[8].text = denda_str
            row_cells[9].text = f"{net:,.1f}"
            
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return output
