import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .templates.header_config import ReportHeaderConfig

class SlipWordExporter:
    def generate(self, summary, harvester_name: str, employee_number: str, period_name: str) -> io.BytesIO:
        document = Document()
        
        header_data = ReportHeaderConfig.get_dynamic_header_data()
        
        # Header
        h1 = document.add_heading(header_data["company_name"], 0)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = document.add_paragraph()
        p.add_run("SLIP GAJI PEMANEN").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Info
        info_table = document.add_table(rows=1, cols=4)
        info_cells = info_table.rows[0].cells
        info_cells[0].text = "Pemanen"
        info_cells[1].text = f": {employee_number} - {harvester_name}"
        info_cells[2].text = "Periode"
        info_cells[3].text = f": {period_name}"
        for cell in info_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    
        document.add_paragraph() # spacing
        
        # Section 1: Rangkuman Prestasi Panen
        p_sec1 = document.add_paragraph()
        r_sec1 = p_sec1.add_run("Rangkuman Prestasi Panen")
        r_sec1.bold = True
        
        table1 = document.add_table(rows=3, cols=2)
        table1.style = 'Table Grid'
        
        r1_cells = table1.rows[0].cells
        r1_cells[0].text = "Total Janjang Masuk"
        r1_cells[1].text = f"{summary.total_valid_bunch_count} jjg"
        
        r2_cells = table1.rows[1].cells
        r2_cells[0].text = "Total Janjang Denda (Mentah, dll)"
        r2_cells[1].text = f"{summary.total_unripe_bunch_count} jjg"
        
        r3_cells = table1.rows[2].cells
        r3_cells[0].text = "Total Tonase Bersih (Netto)"
        r3_cells[1].text = f"{float(summary.total_net_tonnage_kg):,.2f} kg"
        
        document.add_paragraph() # spacing
        
        # Section 2: Rincian Prestasi Harian
        if hasattr(summary, 'daily_records') and summary.daily_records:
            p_sec_daily = document.add_paragraph()
            r_sec_daily = p_sec_daily.add_run("Rincian Prestasi Harian")
            r_sec_daily.bold = True
            
            records_count = len(summary.daily_records)
            table_daily = document.add_table(rows=records_count + 1, cols=6)
            table_daily.style = 'Table Grid'
            
            headers = ["Tanggal", "Jjg Valid", "Jjg Denda", "Netto (kg)", "Premi Brondol (Rp)", "Potongan (Rp)"]
            hdr_cells = table_daily.rows[0].cells
            for idx, header_text in enumerate(headers):
                hdr_cells[idx].text = header_text
                hdr_cells[idx].paragraphs[0].runs[0].bold = True
                
            for i, r in enumerate(summary.daily_records):
                row_cells = table_daily.rows[i + 1].cells
                row_cells[0].text = r.harvest_date.strftime("%d-%m-%Y") if hasattr(r, 'harvest_date') else "-"
                row_cells[1].text = str(r.valid_bunch_count)
                row_cells[2].text = str(r.unripe_bunch_count)
                row_cells[3].text = f"{float(r.net_tonnage_kg):,.2f}"
                row_cells[4].text = f"{float(r.loose_fruit_premium_rupiah):,.0f}"
                row_cells[5].text = f"{float(r.fine_amount_rupiah):,.0f}"
                
            document.add_paragraph() # spacing
        
        # Section 3: Rincian Pendapatan & Potongan
        p_sec2 = document.add_paragraph()
        r_sec2 = p_sec2.add_run("Rincian Pendapatan & Potongan")
        r_sec2.bold = True
        
        tiers_count = len(summary.tier_details) if hasattr(summary, 'tier_details') and summary.tier_details else 0
        table2 = document.add_table(rows=4 + tiers_count, cols=2)
        table2.style = 'Table Grid'
        
        t2_r1 = table2.rows[0].cells
        t2_r1[0].text = "Rincian"
        t2_r1[1].text = "Jumlah (Rp)"
        t2_r1[0].paragraphs[0].runs[0].bold = True
        t2_r1[1].paragraphs[0].runs[0].bold = True
        
        t2_r2 = table2.rows[1].cells
        t2_r2[0].text = "Premi Brondolan (Loose Fruit)"
        t2_r2[1].text = f"{float(summary.total_loose_fruit_premium_rupiah):,.0f}"
        
        current_row = 2
        if tiers_count > 0:
            for t in summary.tier_details:
                row_cells = table2.rows[current_row].cells
                desc = f"Premi Progresif Tier {getattr(t, 'tier_level', 0)} ({float(getattr(t, 'kg_in_tier', 0)):,.2f} kg x Rp {float(getattr(t, 'rate_per_kg', 0)):,.0f})"
                row_cells[0].text = desc
                row_cells[1].text = f"{float(getattr(t, 'subtotal_rupiah', 0)):,.0f}"
                current_row += 1
                
        t2_fine = table2.rows[current_row].cells
        t2_fine[0].text = f"Potongan Denda (Mode: {summary.fine_mode_used})"
        t2_fine[1].text = f"({float(summary.total_fine_rupiah):,.0f})"
        
        current_row += 1
        t2_total = table2.rows[current_row].cells
        t2_total[0].text = "TOTAL PENDAPATAN BERSIH (NET PAY)"
        t2_total[1].text = f"Rp {float(summary.total_net_pay_rupiah):,.0f}"
        t2_total[0].paragraphs[0].runs[0].bold = True
        t2_total[1].paragraphs[0].runs[0].bold = True
        
        # Signatures
        document.add_paragraph()
        document.add_paragraph()
        
        sig_table = document.add_table(rows=1, cols=2)
        sig_cells = sig_table.rows[0].cells
        
        p_left = sig_cells[0].paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_left.add_run("Mengetahui,\n\n\n\n\n(___________________)\nAsisten Divisi / Manager")
        
        p_right = sig_cells[1].paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_right.add_run(f"Diterima Oleh,\n\n\n\n\n(___________________)\n{harvester_name}")
        
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return output
